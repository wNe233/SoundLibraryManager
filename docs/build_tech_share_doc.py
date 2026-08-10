from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "Codex音效管理器开发技术分享.docx"
ARCH_IMG = DOCS / "architecture-flow.png"
TIMELINE_IMG = DOCS / "iteration-timeline.png"
ICON = ROOT / "build" / "icon.png"


INK = "111827"
MUTED = "5F6B7A"
BLUE = "2563EB"
CYAN = "0891B2"
GREEN = "059669"
AMBER = "B45309"
RED = "B91C1C"
LIGHT = "F8FAFC"
BORDER = "DDE3EA"
FILL = "EEF4FF"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for item in candidates:
        try:
            return ImageFont.truetype(item, size=size, index=0)
        except Exception:
            continue
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline=None, radius=22, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, box, text, fnt, fill=INK):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=6, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=fnt, fill=fill, spacing=6, align="center")


def wrap_text(text, max_chars):
    lines = []
    current = ""
    for ch in text:
        current += ch
        if len(current) >= max_chars:
            lines.append(current)
            current = ""
    if current:
        lines.append(current)
    return "\n".join(lines)


def make_architecture_image():
    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = font(42, True)
    h = font(28, True)
    body = font(23)
    small = font(20)

    draw.text((70, 54), "音效管理器技术架构：轻索引 + 波形缓存 + 原生拖拽", font=title, fill=f"#{INK}")
    draw.text((70, 112), "核心原则：不复制原始音频，只建立本地索引；重任务进入缓存，交互层保持轻量。", font=body, fill=f"#{MUTED}")

    boxes = [
        ((80, 220, 380, 365), "导入层", "文件/文件夹拖入\n保留原始层级\n读取基本元数据", BLUE),
        ((480, 220, 780, 365), "分析层", "生成真实波形\n内容哈希去重\n自动识别筛选", CYAN),
        ((880, 220, 1180, 365), "本地库", "library.json\nwave-cache\n用户可改缓存位置", GREEN),
        ((1280, 220, 1580, 365), "编辑器连接", "Finder/Explorer 显示\n右键导入媒体池\n原生文件拖拽", AMBER),
        ((250, 555, 550, 710), "渲染层", "卡片/列表视图\n分批显示\n只重绘可见波形", BLUE),
        ((650, 555, 950, 710), "交互层", "指针即预览\n空格播放\n标签与筛选", CYAN),
        ((1050, 555, 1350, 710), "发布层", "macOS App/DMG\nWindows x64 包\n原生模块构建", RED),
    ]

    for xy, name, desc, color in boxes:
        rounded_box(draw, xy, fill="#FFFFFF", outline=f"#{BORDER}", radius=26, width=3)
        draw.rounded_rectangle((xy[0], xy[1], xy[0] + 12, xy[3]), radius=6, fill=f"#{color}")
        draw.text((xy[0] + 34, xy[1] + 28), name, font=h, fill=f"#{INK}")
        draw.multiline_text((xy[0] + 34, xy[1] + 72), desc, font=small, fill=f"#{MUTED}", spacing=10)

    arrows = [
        ((390, 292), (470, 292)), ((790, 292), (870, 292)), ((1190, 292), (1270, 292)),
        ((1030, 365), (860, 545)), ((780, 365), (790, 545)), ((530, 365), (420, 545)),
        ((550, 632), (640, 632)), ((950, 632), (1040, 632)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#94A3B8", width=5)
        ex, ey = end
        sx, sy = start
        if ex >= sx:
            tri = [(ex, ey), (ex - 16, ey - 10), (ex - 16, ey + 10)]
        else:
            tri = [(ex, ey), (ex + 16, ey - 10), (ex + 16, ey + 10)]
        draw.polygon(tri, fill="#94A3B8")

    rounded_box(draw, (70, 785, 1530, 850), fill="#F8FAFC", outline="#E2E8F0", radius=20, width=2)
    draw.text((100, 805), "从沟通中形成的产品判断：高级感不是堆功能，而是让剪辑师少找、少等、少切换，同时保持可控。", font=body, fill=f"#{INK}")
    img.save(ARCH_IMG)


def make_timeline_image():
    img = Image.new("RGB", (1600, 980), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = font(42, True)
    h = font(25, True)
    small = font(19)
    draw.text((70, 55), "版本迭代路线：从可运行原型到 1.0 发布版", font=title, fill=f"#{INK}")
    draw.text((70, 112), "每个阶段都由真实使用反馈驱动，重点从“能用”逐步转向“顺手、稳定、可分享”。", font=small, fill=f"#{MUTED}")

    stages = [
        ("0.1.0", "基础原型", "Electron 框架、本地导入、深浅色、预览播放、基础波形。", BLUE),
        ("0.1.1-0.1.3", "素材库体验", "文件夹层级、右键菜单、卡片/列表、标签筛选、浅色主题修正。", CYAN),
        ("0.1.4-0.1.8", "拖拽攻坚", "放弃单纯 JS 拖拽，转向系统原生拖拽；逐步打通 macOS、剪映、达芬奇。", AMBER),
        ("0.1.9-0.1.14", "跨平台封装", "Windows 虚拟机验证、原生模块构建脚本、运行日志、打包流程稳定化。", RED),
        ("1.0.0", "发布版收敛", "重复音频识别、收藏夹、缺失检测、缓存设置、大库性能优化、体积清理。", GREEN),
    ]

    x = 120
    y = 230
    gap = 155
    for i, (ver, name, desc, color) in enumerate(stages):
        cy = y + i * gap
        draw.ellipse((88, cy + 20, 128, cy + 60), fill=f"#{color}")
        draw.line((108, cy + 60, 108, cy + gap), fill="#CBD5E1", width=4)
        rounded_box(draw, (170, cy, 1480, cy + 112), fill="#FFFFFF", outline=f"#{BORDER}", radius=24, width=2)
        draw.text((205, cy + 20), ver, font=h, fill=f"#{color}")
        draw.text((390, cy + 20), name, font=h, fill=f"#{INK}")
        draw.text((390, cy + 60), desc, font=small, fill=f"#{MUTED}")
    draw.line((108, y + 4 * gap + 60, 108, y + 4 * gap + 100), fill="#FFFFFF", width=8)

    rounded_box(draw, (70, 900, 1530, 940), fill="#F8FAFC", outline="#E2E8F0", radius=18, width=2)
    draw.text((100, 910), "版本越往后，新增功能越少，修复体验细节、跨平台边界和性能问题越多，这也是工具型产品成熟的标志。", font=small, fill=f"#{INK}")
    img.save(TIMELINE_IMG)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)


def set_run_font(run, east_asia="PingFang SC"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_para_font(paragraph, size=None, color=None, bold=None):
    for run in paragraph.runs:
        set_run_font(run)
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold


def add_para(doc, text="", style=None, size=None, color=None, bold=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if text:
        run = p.add_run(text)
        set_run_font(run)
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    color = BLUE if level <= 2 else "1F4D78"
    set_para_font(p, size=16 if level == 1 else 13 if level == 2 else 11.5, color=color, bold=True)
    return p


def add_callout(doc, title, body, color=FILL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(widths):
        table.columns[i].width = Inches(w)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].width = Inches(widths[i])
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_text(hdr[i], text, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], text)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def build_doc():
    make_architecture_image()
    make_timeline_image()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.86)
    sec.right_margin = Inches(0.86)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    if ICON.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(ICON), width=Inches(0.7))
    title = add_para(doc, "Codex 开发音效管理器技术分享", size=26, color=INK, bold=True, after=4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = add_para(
        doc,
        "从影视剪辑师的真实工作流出发，把一个 Electron 原型打磨成本地化、跨平台、可发布的音效素材库工具",
        size=12.5,
        color=MUTED,
        after=10,
    )
    add_callout(
        doc,
        "一句话总结",
        "这次开发不是单纯堆功能，而是在反复反馈中不断删掉无效设计、强化关键路径：导入、查找、预览、拖拽、整理、发布。",
        "EEF6FF",
    )

    add_heading(doc, "项目背景：从剪辑师的痛点开始", 1)
    add_para(
        doc,
        "这个音效管理器的目标很清楚：不做联网素材平台，不内置音效，不做账号系统，只帮助剪辑师把本地音效管理得更快、更清楚、更接近高级素材库的使用体验。用户最早提出的核心要求是：通过可执行程序直接运行，中文界面，支持深浅色切换，支持实时预览，并尽可能把音效从软件内直接拖入达芬奇或剪映时间线。",
    )
    add_para(
        doc,
        "在后续沟通中，需求逐渐从“能不能实现”变成“怎样才算专业工具”：真实波形、文件夹层级保留、卡片与列表双视图、右键目录管理、自定义标签、重复音频识别、收藏夹、缺失检测、缓存设置、Windows 完整移植，以及大量音频时仍然流畅。",
    )

    add_heading(doc, "产品判断：高级感来自明确、轻巧和稳定", 1)
    add_bullet(doc, "界面要有设计感，但动画不能抢任务。悬停、播放头、下拉、导入进度的动画都服务于状态表达。")
    add_bullet(doc, "素材库必须轻，不复制原始文件，只保存索引、标签、分析结果和波形缓存。")
    add_bullet(doc, "右侧详情页只服务单个音频，相同音频识别这类批处理功能独立成弹窗，避免信息结构混乱。")
    add_bullet(doc, "剪辑师最常用的动作要短：鼠标指到哪里就预览哪里，空格从播放头开始，卡片和列表都能直接操作。")
    add_bullet(doc, "跨平台不是换个打包命令，尤其是文件拖拽到剪辑软件，必须进入系统原生拖拽层。")

    doc.add_picture(str(ARCH_IMG), width=Inches(6.45))
    cap = add_para(doc, "图 1  这个项目的关键架构：前端负责体验，主进程负责文件系统，缓存层负责减轻重复计算，原生模块负责拖拽边界。", size=9, color=MUTED, after=8)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "技术方案：Electron 做壳，原生能力补边界", 1)
    add_heading(doc, "本地库与缓存", 2)
    add_para(
        doc,
        "主进程负责扫描音频文件、读取元数据、保存 library.json，并把真实波形和识别结果写入 wave-cache。缓存位置允许用户调整，避免将来 Windows 用户把 C 盘占满。索引文件只记录路径、层级、标签、收藏、缺失状态、分析结果和缓存引用，不复制原始音频。",
    )
    add_heading(doc, "真实波形与识别筛选", 2)
    add_para(
        doc,
        "渲染进程通过 Web Audio API 解码音频，提取峰值数组绘制真实波形，并用轻量规则识别时间形态、频率特征、动态变化和音高趋势。这里没有追求重型 AI，而是用足够快的启发式规则满足“快速筛选音效”的目标。",
    )
    add_heading(doc, "拖拽到剪辑软件", 2)
    add_para(
        doc,
        "拖拽是最难的部分。早期尝试过 Electron 的 startDrag、DataTransfer、DownloadURL 等方式，但在达芬奇和剪映这类专业软件中并不稳定。最终方案是保留 Electron 作为界面层，同时写 macOS 和 Windows 的原生文件拖拽模块，让系统认为拖出的就是一个真实文件。",
    )

    add_heading(doc, "版本迭代：每一次反馈都改变了产品形态", 1)
    doc.add_picture(str(TIMELINE_IMG), width=Inches(6.45))
    cap = add_para(doc, "图 2  版本演进不是线性加功能，而是围绕剪辑师工作流持续收敛。", size=9, color=MUTED, after=8)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = [
        ("0.1.0", "基础原型", "搭建 Electron 应用；支持本地导入、中文界面、深浅色切换、播放预览、初版波形和右侧详情。"),
        ("0.1.1", "界面与导入修复", "优化波形视觉；去掉无效总时长；修复文件/文件夹拖入；增加卡片模式并保留列表模式。"),
        ("0.1.2-0.1.3", "素材管理", "支持导入文件夹层级、右键新建/重命名/移除目录、框选与多选、标签筛选、浅色主题调整。"),
        ("0.1.4-0.1.8", "拖拽攻坚", "多轮验证 JS 拖拽失败后，改为参考系统行为；最终用原生拖拽让达芬奇和剪映识别真实音频文件。"),
        ("0.1.9-0.1.14", "Windows 移植", "Windows 原生模块、构建脚本、虚拟机测试、运行日志、打包脚本反复修正，解决环境与发布流程问题。"),
        ("1.0.0", "发布收敛", "加入重复音频识别、收藏夹、缺失检测、缓存设置；优化大库性能；清理发布包体积，作为首个发布版。"),
    ]
    add_table(doc, ["版本", "主题", "主要迭代内容"], rows, [0.9, 1.25, 4.15])

    add_heading(doc, "几个关键难点的处理", 1)
    add_heading(doc, "一、拖拽不是前端事件，而是系统数据对象", 2)
    add_para(
        doc,
        "浏览器里的拖拽更像网页内部交互，而剪辑软件需要的是系统级文件拖拽。Windows 侧使用 OLE 的 DoDragDrop 和 CF_HDROP 数据对象，macOS 侧使用 NSDraggingSession 与 file URL。这个转向是项目能真正进入达芬奇、剪映工作流的关键。",
    )
    add_heading(doc, "二、波形要真实，但不能拖慢界面", 2)
    add_para(
        doc,
        "波形既是审美元素，也是效率工具。它必须对应真实文件，方便剪辑师判断长短、爆发点和动态变化。实现上先提取峰值数据，再缓存到本地；UI 只绘制需要展示的波形，而不是每次都重新解码音频。",
    )
    add_heading(doc, "三、重复音频识别不能放在详情页", 2)
    add_para(
        doc,
        "重复识别属于批处理任务，不属于单条音频详情。最终把它做成独立功能按钮和弹窗：按内容哈希分组，每组可预览、看波形、看目录路径，并选择保留哪一个。这样用户能判断“是不是同一个音效”，而不是被系统直接替他删除。",
    )
    add_heading(doc, "四、大库性能要从渲染和数据结构一起处理", 2)
    add_para(
        doc,
        "当库里有大量音频时，卡顿来自多个方向：一次性渲染所有卡片、给每个元素绑定大量事件、每次重绘所有波形、目录数量反复全库扫描、悬停预览也触发保存。1.0 中改为分批渲染、可见波形绘制、素材和目录索引、筛选缓存、搜索防抖，并避免悬停时写入完整库文件。",
    )

    add_heading(doc, "1.0 发布版的性能收敛", 1)
    perf_rows = [
        ("列表/卡片渲染", "从一次性全量渲染改为初始显示一批，滚动接近底部再继续加载。", "大库打开和筛选更快。"),
        ("波形绘制", "只重绘当前可见和必要的波形，避免全列表反复 canvas 绘制。", "滚动与播放头更顺。"),
        ("素材索引", "增加 soundById，选中、预览、收藏、重复弹窗不再反复 find。", "交互响应更稳定。"),
        ("目录索引", "增加 folderById、children、descendants、count 缓存。", "左侧层级数量不再反复全库扫描。"),
        ("状态保存", "localStorage 只存 UI 偏好，完整库写 library.json。", "减少频繁写大 JSON。"),
        ("缺失检测", "启动只做轻量抽查，完整检测由用户主动触发。", "启动更快。"),
    ]
    add_table(doc, ["优化点", "做法", "效果"], perf_rows, [1.3, 3.15, 1.85])

    add_heading(doc, "如何让软件水平更好：这次迭代给出的经验", 1)
    add_number(doc, "先抓最短工作流。剪辑师最常用的是找、听、拖、整理，任何功能都要服务这几步。")
    add_number(doc, "用真实反馈淘汰伪功能。比如混响、人声筛选在这个阶段价值不高，就应删掉；右侧详情页不适合放批处理，也应拆出去。")
    add_number(doc, "审美不能只看静态截图。波形、悬停、播放头、导入遮罩和弹窗状态都要在真实操作中顺滑。")
    add_number(doc, "跨平台要提前识别原生边界。文件拖拽、文件路径、权限、打包和系统菜单都不是纯前端问题。")
    add_number(doc, "发布版要做减法。功能稳定后，重点变成性能、缓存、体积、日志、错误提示和用户可理解的打包流程。")

    add_heading(doc, "后续可以继续打磨的方向", 1)
    add_bullet(doc, "真正的虚拟列表：如果未来素材达到数万条，可以继续从分批渲染升级为窗口化渲染。")
    add_bullet(doc, "更细的音频分析：在不牺牲速度的前提下，加入响度、峰值、BPM 或更可靠的频段分析。")
    add_bullet(doc, "编辑器适配矩阵：分别记录达芬奇、剪映、Premiere、Final Cut 的拖拽和导入表现。")
    add_bullet(doc, "库备份与迁移：在仍然保持本地化的前提下，提供导出索引、重连路径、批量修复缺失文件。")
    add_bullet(doc, "发布体验：macOS 公证、Windows 安装包签名、自动生成测试日志包，减少用户手动排查成本。")

    add_heading(doc, "结语", 1)
    add_para(
        doc,
        "这次开发的价值不只是做出了一个音效管理器，更重要的是形成了一套工具型软件的迭代方法：从真实工作流出发，用反馈不断修正信息架构和技术路线，把看得见的界面体验和看不见的系统能力一起打磨。对于面向剪辑师的本地工具来说，真正的高级感不是复杂，而是可靠、直观、轻快，并且在关键动作上不打断创作。",
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Codex 音效管理器开发技术分享 · 1.0 发布复盘")
        set_run_font(run)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
